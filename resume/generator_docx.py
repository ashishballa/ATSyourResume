import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .generator import _clean_list

ACCENT = RGBColor(0x2c, 0x3e, 0x50)
MUTED = RGBColor(0x7f, 0x8c, 0x8d)


def _add_hr(doc):
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2c3e50')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _section_header(doc, text):
    """Add a section label with a rule underneath."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = ACCENT
    _add_hr(doc)


def _set_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def build_docx(data: dict) -> bytes:
    doc = Document()
    _set_margins(doc)

    # Remove default empty paragraph Word adds
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # ── Header ──────────────────────────────────────────────────────────────
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(data.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = ACCENT

    c = data.get("contact", {}) if isinstance(data.get("contact"), dict) else {}
    def _valid(v):
        return v and isinstance(v, str) and "optional" not in v.lower()
    primary = [v for v in [c.get("email"), c.get("phone"), c.get("location")] if _valid(v)]
    secondary = [v for v in [c.get("linkedin"), c.get("github")] if _valid(v)]
    for line_parts in [primary, secondary]:
        if not line_parts:
            continue
        contact_p = doc.add_paragraph("  |  ".join(line_parts))
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(2)
        for run in contact_p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = MUTED
    if primary or secondary:
        # add a small gap after the last contact line before the HR
        doc.paragraphs[-1].paragraph_format.space_after = Pt(6)

    _add_hr(doc)

    # ── Summary ─────────────────────────────────────────────────────────────
    if summary := data.get("summary"):
        _section_header(doc, "PROFESSIONAL SUMMARY")
        sp = doc.add_paragraph(summary)
        sp.paragraph_format.space_after = Pt(4)
        for run in sp.runs:
            run.font.size = Pt(10)

    # ── Skills ──────────────────────────────────────────────────────────────
    if skills := _clean_list(data.get("skills")):
        _section_header(doc, "SKILLS")
        skp = doc.add_paragraph("  •  ".join(skills))
        skp.paragraph_format.space_after = Pt(4)
        for run in skp.runs:
            run.font.size = Pt(10)

    # ── Experience ──────────────────────────────────────────────────────────
    if experience := data.get("experience"):
        _section_header(doc, "EXPERIENCE")
        for job in experience:
            tp = doc.add_paragraph()
            tp.paragraph_format.space_before = Pt(4)
            tp.paragraph_format.space_after = Pt(1)
            title_run = tp.add_run(job.get("title", ""))
            title_run.bold = True
            title_run.font.size = Pt(10)

            mp = doc.add_paragraph(f"{job.get('company', '')}  |  {job.get('dates', '')}")
            mp.paragraph_format.space_after = Pt(2)
            for run in mp.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = MUTED

            for bullet in job.get("bullets", []):
                bp = doc.add_paragraph(style='List Bullet')
                bp.paragraph_format.space_after = Pt(1)
                bp.paragraph_format.left_indent = Inches(0.15)
                run = bp.add_run(bullet)
                run.font.size = Pt(10)

    # ── Education ───────────────────────────────────────────────────────────
    if education := data.get("education"):
        _section_header(doc, "EDUCATION")
        for e in education:
            dp = doc.add_paragraph()
            dp.paragraph_format.space_before = Pt(4)
            dp.paragraph_format.space_after = Pt(1)
            deg_run = dp.add_run(e.get("degree", ""))
            deg_run.bold = True
            deg_run.font.size = Pt(10)

            _gpa = e.get("gpa", "")
            gpa = f"  |  GPA: {_gpa}" if _gpa and "optional" not in str(_gpa).lower() else ""
            mp = doc.add_paragraph(f"{e.get('school', '')}  |  {e.get('dates', '')}{gpa}")
            mp.paragraph_format.space_after = Pt(2)
            for run in mp.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = MUTED

    # ── Certifications ──────────────────────────────────────────────────────
    if certs := _clean_list(data.get("certifications")):
        _section_header(doc, "CERTIFICATIONS")
        for cert in certs:
            cp = doc.add_paragraph(style='List Bullet')
            cp.paragraph_format.space_after = Pt(1)
            cp.paragraph_format.left_indent = Inches(0.15)
            run = cp.add_run(cert)
            run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
